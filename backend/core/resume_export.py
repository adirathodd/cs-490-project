"""
UC-051: Resume Export Service
Multi-format resume export with theming support
Supports: PDF, DOCX, HTML, Plain Text
"""
import base64
import io
import logging
from typing import Dict, Any, List, Optional
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from jinja2 import Template
from django.utils.html import escape

logger = logging.getLogger(__name__)


class ResumeExportError(Exception):
    """Custom exception for resume export errors"""
    pass


# Theme configurations
THEMES = {
    'professional': {
        'name': 'Professional',
        'description': 'Classic business style with conservative formatting',
        'colors': {
            'primary': (31, 78, 121),    # Dark blue
            'secondary': (100, 100, 100),  # Gray
            'accent': (0, 0, 0),          # Black
        },
        'fonts': {
            'heading': 'Calibri',
            'body': 'Calibri',
        },
        'spacing': 'standard',
    },
    'modern': {
        'name': 'Modern',
        'description': 'Contemporary design with accent colors',
        'colors': {
            'primary': (99, 102, 241),    # Indigo
            'secondary': (71, 85, 105),    # Slate
            'accent': (239, 68, 68),      # Red accent
        },
        'fonts': {
            'heading': 'Arial',
            'body': 'Arial',
        },
        'spacing': 'compact',
    },
    'minimal': {
        'name': 'Minimal',
        'description': 'Clean and simple black and white design',
        'colors': {
            'primary': (0, 0, 0),         # Black
            'secondary': (75, 75, 75),    # Dark gray
            'accent': (0, 0, 0),          # Black
        },
        'fonts': {
            'heading': 'Arial',
            'body': 'Arial',
        },
        'spacing': 'compact',
    },
    'creative': {
        'name': 'Creative',
        'description': 'Bold colors and distinctive formatting',
        'colors': {
            'primary': (16, 185, 129),    # Green
            'secondary': (59, 130, 246),   # Blue
            'accent': (239, 68, 68),      # Red
        },
        'fonts': {
            'heading': 'Calibri',
            'body': 'Calibri',
        },
        'spacing': 'standard',
    },
}


def get_available_themes() -> List[Dict[str, str]]:
    """Get list of available export themes"""
    return [
        {
            'id': theme_id,
            'name': theme_data['name'],
            'description': theme_data['description'],
        }
        for theme_id, theme_data in THEMES.items()
    ]


def _format_date_range(start_date: Optional[date], end_date: Optional[date], is_current: bool = False) -> str:
    """Format date range for display"""
    if not start_date:
        return ''
    
    start_str = start_date.strftime('%b %Y')
    
    if is_current:
        return f"{start_str} - Present"
    elif end_date:
        return f"{start_str} - {end_date.strftime('%b %Y')}"
    else:
        return start_str


def collect_profile_data(profile) -> Dict[str, Any]:
    """
    Collect comprehensive profile data for export
    
    Args:
        profile: CandidateProfile instance
        
    Returns:
        Dictionary with structured profile data
    """
    from core.models import CandidateSkill, WorkExperience, Education, Certification, Project
    
    # Basic info
    data = {
        'name': profile.get_full_name() or 'Unknown',
        'email': profile.user.email,
        'phone': profile.phone or '',
        'location': profile.get_full_location() or '',
        'headline': profile.headline or '',
        'summary': profile.summary or '',
        'portfolio_url': profile.portfolio_url or '',
    }
    
    # Skills
    skills = CandidateSkill.objects.filter(candidate=profile).select_related('skill').order_by('order')
    skills_by_category = {}
    for cs in skills:
        category = cs.skill.category or 'Other'
        if category not in skills_by_category:
            skills_by_category[category] = []
        skills_by_category[category].append({
            'name': cs.skill.name,
            'level': cs.level,
            'years': float(cs.years) if cs.years else 0,
        })
    data['skills'] = skills_by_category
    
    # Work Experience
    experiences = []
    for exp in WorkExperience.objects.filter(candidate=profile).order_by('-start_date'):
        exp_data = {
            'company_name': exp.company_name,
            'job_title': exp.job_title,
            'location': exp.location,
            'date_range': _format_date_range(exp.start_date, exp.end_date, exp.is_current),
            'description': exp.description,
            'achievements': exp.achievements if isinstance(exp.achievements, list) else [],
            'skills': [s.name for s in exp.skills_used.all()],
        }
        experiences.append(exp_data)
    data['experiences'] = experiences
    
    # Education
    education = []
    for edu in Education.objects.filter(candidate=profile).order_by('-end_date'):
        edu_data = {
            'institution': edu.institution,
            'degree_type': edu.get_degree_type_display(),
            'field_of_study': edu.field_of_study,
            'date_range': _format_date_range(edu.start_date, edu.end_date, edu.currently_enrolled),
            'gpa': float(edu.gpa) if edu.gpa and not edu.gpa_private else None,
            'honors': edu.honors,
            'achievements': edu.achievements,
        }
        education.append(edu_data)
    data['education'] = education
    
    # Certifications
    certifications = []
    for cert in Certification.objects.filter(candidate=profile).order_by('-issue_date'):
        cert_data = {
            'name': cert.name,
            'issuing_organization': cert.issuing_organization,
            'issue_date': cert.issue_date.strftime('%b %Y'),
            'expiry_date': cert.expiry_date.strftime('%b %Y') if cert.expiry_date and not cert.never_expires else None,
            'credential_id': cert.credential_id,
            'credential_url': cert.credential_url,
        }
        certifications.append(cert_data)
    data['certifications'] = certifications
    
    # Projects
    projects = []
    for proj in Project.objects.filter(candidate=profile).order_by('display_order', '-start_date'):
        proj_data = {
            'name': proj.name,
            'role': proj.role,
            'description': proj.description,
            'date_range': _format_date_range(proj.start_date, proj.end_date),
            'project_url': proj.project_url,
            'outcomes': proj.outcomes,
            'skills': [s.name for s in proj.skills_used.all()],
        }
        projects.append(proj_data)
    data['projects'] = projects
    
    return data


# ============================================================================
# PLAIN TEXT EXPORT
# ============================================================================

def export_plain_text(profile_data: Dict[str, Any]) -> str:
    """
    Export resume as plain text
    Suitable for online application forms and ATS systems
    """
    lines = []
    
    # Header
    lines.append(profile_data['name'].upper())
    lines.append('=' * len(profile_data['name']))
    
    contact_info = []
    if profile_data['email']:
        contact_info.append(profile_data['email'])
    if profile_data['phone']:
        contact_info.append(profile_data['phone'])
    if profile_data['location']:
        contact_info.append(profile_data['location'])
    
    if contact_info:
        lines.append(' | '.join(contact_info))
    
    if profile_data['portfolio_url']:
        lines.append(f"Portfolio: {profile_data['portfolio_url']}")
    
    lines.append('')
    
    # Headline
    if profile_data['headline']:
        lines.append(profile_data['headline'])
        lines.append('')
    
    # Summary
    if profile_data['summary']:
        lines.append('PROFESSIONAL SUMMARY')
        lines.append('-' * 20)
        lines.append(profile_data['summary'])
        lines.append('')
    
    # Skills
    if profile_data['skills']:
        lines.append('SKILLS')
        lines.append('-' * 20)
        for category, skills in profile_data['skills'].items():
            # Handle both string lists and dict lists
            if skills and isinstance(skills[0], dict):
                skill_names = [s['name'] for s in skills]
            else:
                skill_names = skills
            lines.append(f"{category}: {', '.join(skill_names)}")
        lines.append('')
    
    # Experience
    if profile_data['experiences']:
        lines.append('PROFESSIONAL EXPERIENCE')
        lines.append('-' * 20)
        for exp in profile_data['experiences']:
            lines.append(f"{exp['job_title']}")
            lines.append(f"{exp['company_name']} | {exp.get('date_range', '')}")
            if exp.get('location'):
                lines.append(f"Location: {exp['location']}")
            if exp.get('description'):
                lines.append('')
                lines.append(exp['description'])
            if exp.get('achievements'):
                lines.append('')
                for achievement in exp['achievements']:
                    lines.append(f"  * {achievement}")
            lines.append('')
    
    # Education
    if profile_data['education']:
        lines.append('EDUCATION')
        lines.append('-' * 20)
        for edu in profile_data['education']:
            # Handle both database format and extracted format
            if 'degree' in edu:
                # Extracted from LaTeX - already combined
                lines.append(edu['degree'])
            else:
                # Database format
                lines.append(f"{edu['degree_type']} in {edu['field_of_study']}")
            lines.append(f"{edu['institution']} | {edu.get('date_range', '')}")
            if edu.get('gpa'):
                lines.append(f"GPA: {edu['gpa']:.2f}")
            if edu.get('honors'):
                lines.append(f"Honors: {edu['honors']}")
            lines.append('')
    
    # Certifications
    if profile_data['certifications']:
        lines.append('CERTIFICATIONS')
        lines.append('-' * 20)
        for cert in profile_data['certifications']:
            lines.append(f"{cert['name']} - {cert['issuing_organization']}")
            date_info = cert['issue_date']
            if cert['expiry_date']:
                date_info += f" - {cert['expiry_date']}"
            lines.append(date_info)
            if cert['credential_id']:
                lines.append(f"Credential ID: {cert['credential_id']}")
            lines.append('')
    
    # Projects
    if profile_data['projects']:
        lines.append('PROJECTS')
        lines.append('-' * 20)
        for proj in profile_data['projects']:
            lines.append(f"{proj['name']}")
            if proj.get('role'):
                lines.append(f"Role: {proj['role']}")
            if proj.get('date_range'):
                lines.append(f"Duration: {proj['date_range']}")
            if proj.get('description'):
                lines.append('')
                lines.append(proj['description'])
            if proj.get('project_url'):
                lines.append(f"URL: {proj['project_url']}")
            lines.append('')
    
    return '\n'.join(lines)


# ============================================================================
# HTML EXPORT
# ============================================================================

HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ name }} - Resume</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: {{ theme.fonts.body }}, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            line-height: 1.5;
            color: #333;
            max-width: 850px;
            margin: 0 auto;
            padding: 30px 20px;
            background: #fff;
            position: relative;
            font-size: 11pt;
        }
        
        {% if watermark %}
        .watermark {
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 80px;
            font-weight: bold;
            color: rgba(220, 53, 69, 0.15);
            z-index: 0;
            pointer-events: none;
            user-select: none;
        }
        
        .content-wrapper {
            position: relative;
            z-index: 1;
        }
        {% endif %}
        
        .header {
            text-align: center;
            margin-bottom: 20px;
            padding-bottom: 15px;
            border-bottom: 2px solid rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
        }
        
        .name {
            font-family: {{ theme.fonts.heading }}, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            font-size: 28px;
            font-weight: 700;
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            margin-bottom: 8px;
            text-transform: uppercase;
        }
        
        .headline {
            font-size: 14px;
            color: rgb({{ theme.colors.secondary[0] }}, {{ theme.colors.secondary[1] }}, {{ theme.colors.secondary[2] }});
            margin-bottom: 8px;
        }
        
        .contact-info {
            font-size: 11px;
            color: #666;
        }
        
        .contact-info a {
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            text-decoration: none;
        }
        
        .section {
            margin-bottom: 20px;
        }
        
        .section-title {
            font-family: {{ theme.fonts.heading }}, -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
            font-size: 14px;
            font-weight: 700;
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            margin-bottom: 10px;
            padding-bottom: 4px;
            border-bottom: 1px solid rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            text-transform: uppercase;
        }
        
        .summary {
            line-height: 1.6;
            color: #444;
            font-size: 10.5pt;
        }
        
        .experience-item, .education-item, .cert-item, .project-item {
            margin-bottom: 15px;
        }
        
        .item-header {
            display: flex;
            justify-content: space-between;
            align-items: baseline;
            margin-bottom: 5px;
        }
        
        .item-title {
            font-size: 18px;
            font-weight: 600;
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
        }
        
        .item-date {
            font-size: 14px;
            color: #666;
            white-space: nowrap;
        }
        
        .item-subtitle {
            font-size: 16px;
            color: rgb({{ theme.colors.secondary[0] }}, {{ theme.colors.secondary[1] }}, {{ theme.colors.secondary[2] }});
            margin-bottom: 5px;
        }
        
        .item-location {
            font-size: 14px;
            color: #666;
            margin-bottom: 10px;
        }
        
        .achievements {
            list-style-position: outside;
            margin-left: 20px;
        }
        
        .achievements li {
            margin-bottom: 8px;
            line-height: 1.6;
        }
        
        .skills-grid {
            display: grid;
            grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
            gap: 15px;
        }
        
        .skill-category {
            margin-bottom: 10px;
        }
        
        .skill-category-name {
            font-weight: 600;
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            margin-bottom: 5px;
        }
        
        .skill-tags {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
        }
        
        .skill-tag {
            display: inline-block;
            padding: 4px 12px;
            background: rgba({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }}, 0.1);
            color: rgb({{ theme.colors.primary[0] }}, {{ theme.colors.primary[1] }}, {{ theme.colors.primary[2] }});
            border-radius: 12px;
            font-size: 13px;
        }
        
        @media print {
            body {
                padding: 0;
                background: #fff;
            }
            
            .section {
                page-break-inside: avoid;
            }
        }
        
        @media (max-width: 768px) {
            .item-header {
                flex-direction: column;
                align-items: flex-start;
            }
            
            .skills-grid {
                grid-template-columns: 1fr;
            }
        }
    </style>
</head>
<body>
    {% if watermark %}
    <div class="watermark">{{ watermark }}</div>
    <div class="content-wrapper">
    {% endif %}
    
    <div class="header">
        <h1 class="name">{{ name }}</h1>
        {% if headline %}
        <div class="headline">{{ headline }}</div>
        {% endif %}
        <div class="contact-info">
            {% if email %}<a href="mailto:{{ email }}">{{ email }}</a>{% endif %}
            {% if phone %}{% if email %} | {% endif %}{{ phone }}{% endif %}
            {% if location %}{% if email or phone %} | {% endif %}{{ location }}{% endif %}
            {% if portfolio_url %}{% if email or phone or location %}<br>{% endif %}<a href="{{ portfolio_url }}" target="_blank">{{ portfolio_url }}</a>{% endif %}
        </div>
    </div>
    
    {% if summary %}
    <div class="section">
        <h2 class="section-title">Professional Summary</h2>
        <div class="summary">{{ summary }}</div>
    </div>
    {% endif %}
    
    {% if skills %}
    <div class="section">
        <h2 class="section-title">Skills</h2>
        <div class="skills-grid">
            {% for category, skill_list in skills.items() %}
            <div class="skill-category">
                <div class="skill-category-name">{{ category }}</div>
                <div class="skill-tags">
                    {% for skill in skill_list %}
                    <span class="skill-tag">{{ skill.name }}</span>
                    {% endfor %}
                </div>
            </div>
            {% endfor %}
        </div>
    </div>
    {% endif %}
    
    {% if experiences %}
    <div class="section">
        <h2 class="section-title">Professional Experience</h2>
        {% for exp in experiences %}
        <div class="experience-item">
            <div class="item-header">
                <div class="item-title">{{ exp.job_title }}</div>
                <div class="item-date">{{ exp.date_range }}</div>
            </div>
            <div class="item-subtitle">{{ exp.company_name }}</div>
            {% if exp.location %}
            <div class="item-location">{{ exp.location }}</div>
            {% endif %}
            {% if exp.description %}
            <p>{{ exp.description }}</p>
            {% endif %}
            {% if exp.achievements %}
            <ul class="achievements">
                {% for achievement in exp.achievements %}
                <li>{{ achievement }}</li>
                {% endfor %}
            </ul>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if education %}
    <div class="section">
        <h2 class="section-title">Education</h2>
        {% for edu in education %}
        <div class="education-item">
            <div class="item-header">
                <div class="item-title">{{ edu.degree_type }}{% if edu.field_of_study %} in {{ edu.field_of_study }}{% endif %}</div>
                <div class="item-date">{{ edu.date_range }}</div>
            </div>
            <div class="item-subtitle">{{ edu.institution }}</div>
            {% if edu.gpa %}
            <div class="item-location">GPA: {{ "%.2f"|format(edu.gpa) }}</div>
            {% endif %}
            {% if edu.honors %}
            <p><strong>Honors:</strong> {{ edu.honors }}</p>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if certifications %}
    <div class="section">
        <h2 class="section-title">Certifications</h2>
        {% for cert in certifications %}
        <div class="cert-item">
            <div class="item-title">{{ cert.name }}</div>
            <div class="item-subtitle">{{ cert.issuing_organization }}</div>
            <div class="item-location">
                Issued: {{ cert.issue_date }}
                {% if cert.expiry_date %} | Expires: {{ cert.expiry_date }}{% endif %}
            </div>
            {% if cert.credential_url %}
            <a href="{{ cert.credential_url }}" target="_blank">View Credential</a>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if projects %}
    <div class="section">
        <h2 class="section-title">Projects</h2>
        {% for proj in projects %}
        <div class="project-item">
            <div class="item-header">
                <div class="item-title">{{ proj.name }}</div>
                {% if proj.date_range %}
                <div class="item-date">{{ proj.date_range }}</div>
                {% endif %}
            </div>
            {% if proj.role %}
            <div class="item-subtitle">{{ proj.role }}</div>
            {% endif %}
            {% if proj.description %}
            <p>{{ proj.description }}</p>
            {% endif %}
            {% if proj.project_url %}
            <a href="{{ proj.project_url }}" target="_blank">View Project</a>
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if watermark %}
    </div>
    {% endif %}
</body>
</html>
"""


def export_html(profile_data: Dict[str, Any], theme: str = 'professional', watermark: str = '') -> str:
    """
    Export resume as HTML
    Suitable for web portfolios and online viewing
    
    Args:
        profile_data: Profile data dictionary
        theme: Theme ID (professional, modern, minimal, creative)
        watermark: Optional watermark text
        
    Returns:
        HTML string
    """
    theme_config = THEMES.get(theme, THEMES['professional'])
    
    # Ensure user-provided text is escaped to prevent XSS in generated HTML.
    def _escape_value(v):
        if isinstance(v, str):
            return escape(v)
        if isinstance(v, dict):
            return {k: _escape_value(val) for k, val in v.items()}
        if isinstance(v, list):
            return [_escape_value(i) for i in v]
        return v

    safe_profile = _escape_value(profile_data)

    template = Template(HTML_TEMPLATE)
    html = template.render(
        **safe_profile,
        theme=theme_config,
        watermark=escape(watermark) if watermark else '',
    )
    
    return html


# ============================================================================
# DOCX (Word) EXPORT
# ============================================================================

def export_docx(profile_data: Dict[str, Any], theme: str = 'professional', watermark: str = '') -> bytes:
    """
    Export resume as Word document (.docx)
    Formatted to match Jake's Resume template styling
    
    Args:
        profile_data: Profile data dictionary
        theme: Theme ID (professional, modern, minimal, creative)
        watermark: Optional watermark text
        
    Returns:
        DOCX file as bytes
    """
    theme_config = THEMES.get(theme, THEMES['professional'])
    colors = theme_config['colors']
    
    doc = Document()
    
    # Set narrow margins to match Jake's template
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = theme_config['fonts']['body']
    font.size = Pt(10)  # Match Jake's template smaller body text
    
    # Header - Name (larger, centered, bold)
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile_data['name'].upper())
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(*colors['primary'])
    name_para.paragraph_format.space_after = Pt(2)
    
    # Headline
    if profile_data['headline']:
        headline_para = doc.add_paragraph()
        headline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline_run = headline_para.add_run(profile_data['headline'])
        headline_run.font.size = Pt(11)
        headline_run.font.color.rgb = RGBColor(*colors['secondary'])
        headline_para.paragraph_format.space_after = Pt(2)
    
    # Contact info - condensed on one line
    contact_parts = []
    if profile_data['email']:
        contact_parts.append(profile_data['email'])
    if profile_data['phone']:
        contact_parts.append(profile_data['phone'])
    if profile_data['location']:
        contact_parts.append(profile_data['location'])
    
    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(' | '.join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)
        contact_para.paragraph_format.space_after = Pt(2)
    
    if profile_data['portfolio_url']:
        url_para = doc.add_paragraph()
        url_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        url_run = url_para.add_run(profile_data['portfolio_url'])
        url_run.font.size = Pt(9)
        url_run.font.color.rgb = RGBColor(*colors['primary'])
        url_para.paragraph_format.space_after = Pt(8)
    
    # Add horizontal line after header
    doc.add_paragraph('_' * 80)
    
    # Helper function to add section heading (Jake's template style: UPPERCASE, underlined)
    def add_section_heading(text):
        para = doc.add_paragraph()
        run = para.add_run(text.upper())
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*colors['primary'])
        run.underline = True
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
    
    # Summary
    if profile_data['summary']:
        add_section_heading('Summary')
        summary_para = doc.add_paragraph(profile_data['summary'])
        summary_para.paragraph_format.space_after = Pt(8)
    
    # Skills - condensed format
    if profile_data['skills']:
        add_section_heading('Technical Skills')
        for category, skills in profile_data['skills'].items():
            skill_para = doc.add_paragraph()
            cat_run = skill_para.add_run(f"{category}: ")
            cat_run.bold = True
            cat_run.font.size = Pt(10)
            # Handle both string lists and dict lists
            if skills and isinstance(skills[0], dict):
                skill_names = ', '.join([s['name'] for s in skills])
            else:
                skill_names = ', '.join(skills)
            skill_run = skill_para.add_run(skill_names)
            skill_run.font.size = Pt(10)
            skill_para.paragraph_format.space_after = Pt(2)
    
    # Experience - Jake's template format
    if profile_data['experiences']:
        add_section_heading('Experience')
        for i, exp in enumerate(profile_data['experiences']):
            # Create table for better alignment (Jake's template style)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Left cell: Job title (bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(exp['job_title'])
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Date range (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if exp['date_range']:
                right_run = right_para.add_run(exp['date_range'])
                right_run.font.size = Pt(10)
            
            # Remove table borders for clean look
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
            
            # Company and location line (italic)
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(exp['company_name'])
            company_run.italic = True
            company_run.font.size = Pt(10)
            if exp['location']:
                location_run = company_para.add_run(f" | {exp['location']}")
                location_run.italic = True
                location_run.font.size = Pt(10)
            company_para.paragraph_format.space_after = Pt(2)
            
            # Description
            if exp['description']:
                desc_para = doc.add_paragraph(exp['description'])
                desc_para.paragraph_format.space_after = Pt(2)
            
            # Achievements (bullet points)
            if exp['achievements']:
                for achievement in exp['achievements']:
                    bullet_para = doc.add_paragraph(achievement, style='List Bullet')
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
                    bullet_para.paragraph_format.space_after = Pt(2)
            
            # Add spacing between experiences (except last)
            if i < len(profile_data['experiences']) - 1:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Education - Jake's template format
    if profile_data['education']:
        add_section_heading('Education')
        for i, edu in enumerate(profile_data['education']):
            # Create table for alignment
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Left cell: Institution (bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(edu['institution'])
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Date range (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if edu.get('date_range'):
                right_run = right_para.add_run(edu['date_range'])
                right_run.font.size = Pt(10)
            
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
            
            # Degree line (italic)
            degree_para = doc.add_paragraph()
            if 'degree' in edu:
                degree_text = edu['degree']
            else:
                degree_text = edu.get('degree_type', '')
                if edu.get('field_of_study'):
                    degree_text += f" in {edu['field_of_study']}"
            degree_run = degree_para.add_run(degree_text)
            degree_run.italic = True
            degree_run.font.size = Pt(10)
            degree_para.paragraph_format.space_after = Pt(2)
            
            # GPA and honors
            gpa = edu.get('gpa')
            honors = edu.get('honors')
            if gpa or honors:
                details = []
                if gpa:
                    details.append(f"GPA: {gpa:.2f}")
                if honors:
                    details.append(f"Honors: {honors}")
                details_para = doc.add_paragraph(' | '.join(details))
                details_para.paragraph_format.left_indent = Inches(0.25)
                details_para.paragraph_format.space_after = Pt(2)
            
            # Add spacing between education entries (except last)
            if i < len(profile_data['education']) - 1:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Certifications - Jake's template format
    if profile_data['certifications']:
        add_section_heading('Certifications')
        for i, cert in enumerate(profile_data['certifications']):
            # Create table for alignment
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Left cell: Cert name (bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(cert['name'])
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Issue date (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_info = f"Issued: {cert['issue_date']}"
            if cert['expiry_date']:
                date_info += f" | Expires: {cert['expiry_date']}"
            right_run = right_para.add_run(date_info)
            right_run.font.size = Pt(10)
            
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
            
            # Organization (italic)
            org_para = doc.add_paragraph()
            org_run = org_para.add_run(cert['issuing_organization'])
            org_run.italic = True
            org_run.font.size = Pt(10)
            org_para.paragraph_format.space_after = Pt(2)
            
            # Add spacing between certifications (except last)
            if i < len(profile_data['certifications']) - 1:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Projects - Jake's template format
    if profile_data['projects']:
        add_section_heading('Projects')
        for i, proj in enumerate(profile_data['projects']):
            # Create table for alignment
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.allow_autofit = False
            
            # Left cell: Project name (bold)
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            left_run = left_para.add_run(proj['name'])
            left_run.bold = True
            left_run.font.size = Pt(11)
            
            # Right cell: Date range (aligned right)
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if proj.get('date_range'):
                right_run = right_para.add_run(proj['date_range'])
                right_run.font.size = Pt(10)
            
            # Remove table borders
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
            
            # Role (italic)
            role = proj.get('role')
            if role:
                role_para = doc.add_paragraph()
                role_run = role_para.add_run(role)
                role_run.italic = True
                role_run.font.size = Pt(10)
                role_para.paragraph_format.space_after = Pt(2)
            
            # Description
            if proj.get('description'):
                desc_para = doc.add_paragraph(proj['description'])
                desc_para.paragraph_format.left_indent = Inches(0.25)
                desc_para.paragraph_format.space_after = Pt(2)
            
            # URL
            if proj.get('project_url'):
                url_para = doc.add_paragraph(proj['project_url'])
                url_para.paragraph_format.left_indent = Inches(0.25)
                url_para.runs[0].font.color.rgb = RGBColor(*colors['primary'])
                url_para.runs[0].font.size = Pt(9)
                url_para.paragraph_format.space_after = Pt(2)
            
            # Add spacing between projects (except last)
            if i < len(profile_data['projects']) - 1:
                doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Add watermark if provided - make it prominent and visible
    if watermark:
        section = doc.sections[0]
        
        # Add to header
        header = section.header
        watermark_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        watermark_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        watermark_run = watermark_para.add_run(watermark.upper())
        watermark_run.font.size = Pt(16)
        watermark_run.font.bold = True
        watermark_run.font.color.rgb = RGBColor(220, 53, 69)  # Red color
        
        # Add to footer for visibility on all pages
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(watermark.upper())
        footer_run.font.size = Pt(16)
        footer_run.font.bold = True
        footer_run.font.color.rgb = RGBColor(220, 53, 69)  # Red color
    
    # Save to bytes
    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    
    return docx_io.read()


# ============================================================================
# MAIN EXPORT FUNCTION
# ============================================================================

def export_resume(
    profile=None,
    format_type: str = '',
    theme: str = 'professional',
    watermark: str = '',
    filename: Optional[str] = None,
    profile_data: Optional[Dict[str, Any]] = None
) -> Dict[str, Any]:
    """
    Main resume export function
    
    Args:
        profile: CandidateProfile instance (optional if profile_data provided)
        format_type: Export format ('pdf', 'docx', 'html', 'txt')
        theme: Theme ID
        watermark: Optional watermark text
        filename: Optional custom filename (without extension)
        profile_data: Optional pre-formatted profile data dict (for AI-generated resumes)
        
    Returns:
        Dictionary with export results:
        {
            'content': bytes or string,
            'content_type': MIME type,
            'filename': suggested filename,
            'format': format type,
        }
    """
    # Validate format
    valid_formats = ['pdf', 'docx', 'html', 'txt']
    if format_type not in valid_formats:
        raise ResumeExportError(f"Invalid format: {format_type}. Must be one of {valid_formats}")
    
    # Validate theme
    if theme not in THEMES:
        logger.warning(f"Invalid theme: {theme}. Using 'professional'")
        theme = 'professional'
    
    # Collect profile data - use provided data or collect from profile
    if profile_data is None:
        if profile is None:
            raise ResumeExportError("Either profile or profile_data must be provided")
        profile_data = collect_profile_data(profile)
    
    # Generate filename if not provided
    if not filename:
        name_parts = profile_data.get('name', 'Resume').replace(' ', '_')
        filename = f"{name_parts}_Resume"
    
    # Export based on format
    if format_type == 'txt':
        content = export_plain_text(profile_data)
        return {
            'content': content,
            'content_type': 'text/plain',
            'filename': f"{filename}.txt",
            'format': 'txt',
        }
    
    elif format_type == 'html':
        content = export_html(profile_data, theme, watermark)
        return {
            'content': content,
            'content_type': 'text/html',
            'filename': f"{filename}.html",
            'format': 'html',
        }
    
    elif format_type == 'docx':
        content = export_docx(profile_data, theme, watermark)
        return {
            'content': content,
            'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'filename': f"{filename}.docx",
            'format': 'docx',
        }
    
    elif format_type == 'pdf':
        # For PDF, we'll use the existing LaTeX compilation from resume_ai
        # This is a placeholder - in production, you might want to convert HTML to PDF
        # or use the existing LaTeX system
        raise ResumeExportError("PDF export requires LaTeX compilation. Use the AI resume generator for PDF export.")
    
    raise ResumeExportError(f"Export format {format_type} not implemented")
