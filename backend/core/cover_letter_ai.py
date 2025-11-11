"""
Gemini-powered cover letter content generation for UC-056.

This module reuses candidate/job snapshot helpers from resume_ai, enriches with
company research context, and generates multiple structured cover letter
variations in JSON via Gemini's API.
"""
from __future__ import annotations

import json
import logging
import re
from typing import Any, Dict, List, Sequence

from django.conf import settings
from django.utils import timezone
from django.utils.text import slugify

from core.models import CandidateProfile, JobEntry, Company, CompanyResearch
from core import resume_ai

logger = logging.getLogger(__name__)

# Soft limits to keep prompts suitable for free-tier generation
MAX_ACHIEVEMENTS = 12
MAX_NEWS = 3
MAX_PARAGRAPHS = 3
MAX_VARIATIONS = 3


class CoverLetterAIError(Exception):
    """Raised when we cannot return AI-generated cover letter content."""


TONE_STYLES = {
    # Backwards-compatible plus UC-058 tones
    'formal': 'Polished, concise, and formal with clear impact statements.',
    'professional': 'Polished, concise, and formal with clear impact statements.',
    'casual': 'Friendly, conversational, and approachable while professional.',
    'enthusiastic': 'Energetic, positive, and motivational with clear excitement.',
    'analytical': 'Data-driven, evidence-focused, and logically structured.',
    'warm': 'Approachable, collaborative, people-first tone with empathy.',
    'innovative': 'Forward-looking, curious, and product/impact oriented.',
    'customer_centric': 'Customer-obsessed voice focusing on outcomes and value.',
    'data_driven': 'Analytical tone weaving metrics and experimentation.',
    'concise': 'Short, crisp sentences with no fluff.',
    'balanced': 'Professional tone blending warmth and measurable results.',
}


def _dedupe(seq: Sequence[str]) -> List[str]:
    seen = set(); out = []
    for item in seq or []:
        if not item:
            continue
        lowered = str(item).strip().lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        out.append(str(item).strip())
    return out


def _collect_achievements(candidate_snapshot: Dict[str, Any], limit: int = MAX_ACHIEVEMENTS) -> List[str]:
    achievements: List[str] = []
    for exp in candidate_snapshot.get('experiences', []) or []:
        for a in (exp.get('achievements') or []):
            txt = (a or '').strip()
            if txt:
                achievements.append(txt)
            if len(achievements) >= limit:
                break
        if len(achievements) >= limit:
            break
    return achievements[:limit]


def build_company_research_snapshot(company_name: str | None) -> Dict[str, Any]:
    """Return compact research info for a company if available; create stubs if missing.

    Structure:
    {
      "company_name": "Acme",
      "culture_keywords": ["ownership", "innovation"],
      "recent_news": [{"title": "...", "url": "...", "date": "YYYY-MM-DD"}],
      "mission_statement": "..."
    }
    """
    if not company_name:
        return {"company_name": "", "culture_keywords": [], "recent_news": [], "mission_statement": ""}
    try:
        company = Company.objects.filter(name__iexact=company_name).first()
        if not company:
            # Create with minimal data to allow later enrichment
            domain = company_name.lower().replace(' ', '').replace(',', '').replace('.', '') + '.com'
            company = Company.objects.create(name=company_name, domain=domain)
            CompanyResearch.objects.create(company=company)
        # Try to fetch research
        research = getattr(company, 'research', None)
        snapshot = {
            'company_name': company.name,
            'culture_keywords': list(research.culture_keywords) if getattr(research, 'culture_keywords', None) else [],
            'recent_news': list(research.recent_news)[:MAX_NEWS] if getattr(research, 'recent_news', None) else [],
            'mission_statement': getattr(research, 'mission_statement', '') or '',
            'description': getattr(research, 'description', '') or '',
        }
        return snapshot
    except Exception as e:
        logger.warning("Failed to build company research snapshot for %s: %s", company_name, e)
        return {"company_name": company_name, "culture_keywords": [], "recent_news": [], "mission_statement": ""}


def build_generation_prompt(
    candidate_snapshot: Dict[str, Any],
    job_snapshot: Dict[str, Any],
    research_snapshot: Dict[str, Any],
    tone: str,
    variation_count: int,
    *,
    length: str | None = None,
    writing_style: str | None = None,
    company_culture: str | None = None,
    industry: str | None = None,
    custom_instructions: str | None = None,
) -> str:
    tone_descriptor = TONE_STYLES.get(tone, TONE_STYLES['balanced'])
    candidate_json = json.dumps(candidate_snapshot, indent=2, ensure_ascii=False)
    job_json = json.dumps(job_snapshot, indent=2, ensure_ascii=False)
    research_json = json.dumps(research_snapshot, indent=2, ensure_ascii=False)

    schema = """
{
    "shared_analysis": {
        "personalization_strategy": "How to tailor to role and company culture.",
        "key_achievements": ["List of most relevant achievements from candidate"],
        "news_to_reference": ["Optional news headline strings used"],
        "tone_rationale": "Why this tone matches the culture"
    },
    "variations": [
        {
            "label": "Warm and Data-Driven",
            "tone": "warm",
            "opening_paragraph": "Personalized intro referencing company, role, and motivation.",
            "body_paragraphs": [
                "Paragraph focused on relevant experience and quantified results.",
                "Optional second para tying achievements to job requirements and culture."
            ],
            "closing_paragraph": "Confident, polite CTA with availability and thanks.",
            "achievements_referenced": ["achievement 1", "achievement 2"],
            "keywords_used": ["Python", "APIs", "Cloud"],
            "news_citations": [
                {"title": "Acme announces new product", "url": "https://...", "date": "2025-10-10"}
            ]
        }
    ]
}
""".strip()

    rules = f"""
You are ResumeRocket AI. Generate {variation_count} cover letter variations as JSON ONLY (no markdown fences).

Tone style: {tone} → {tone_descriptor}

Additional user preferences:
- Desired length: {length or 'standard'}
- Writing style: {writing_style or 'balanced (direct narrative)'}
- Company culture match: {company_culture or 'auto-detect (startup/corporate)'}
- Industry guidance: {industry or 'none'}
{f'- Custom instructions: {custom_instructions}' if custom_instructions else ''}

STRICT RULES:
- Use ONLY facts from candidate_snapshot, job_snapshot, and research_snapshot. Do NOT invent companies, dates, titles, or achievements.
- Refer to the company and role EXACTLY as provided.
- Opening paragraph must be personalized to the company/role and can briefly reference mission or recent news if available.
- Body paragraphs must highlight specific achievements with metrics where possible and map to job requirements.
- Closing paragraph must include a polite call-to-action and availability.
- Match tone cues to culture_keywords when provided.
- Keep each paragraph under 140 words. Use professional writing style.
- Provide 1-2 body paragraphs total.
- Output valid JSON matching the schema below. No comments or extra text.

TONE CONSISTENCY VALIDATION:
- Ensure the selected tone is applied consistently across opening, body and closing paragraphs. If you must soften or vary tone for readability, explain the choice in shared_analysis.tone_rationale.

Schema:
{schema}

candidate_snapshot:
{candidate_json}

job_snapshot:
{job_json}

research_snapshot:
{research_json}
""".strip()
    return rules


def _strip_code_fence(text: str) -> str:
    s = text.strip()
    if s.startswith('```'):
        s = re.sub(r'^```(?:json)?', '', s, count=1, flags=re.IGNORECASE).strip()
        if s.endswith('```'):
            s = s[:-3].strip()
    return s


def parse_payload(raw_text: str) -> Dict[str, Any]:
    try:
        return json.loads(_strip_code_fence(raw_text))
    except json.JSONDecodeError as exc:
        logger.error('Failed to parse Gemini cover letter JSON: %s', exc)
        raise CoverLetterAIError('Gemini returned an unreadable response.') from exc


def _normalize_variations(
    parsed: Dict[str, Any],
    candidate_snapshot: Dict[str, Any],
    job_snapshot: Dict[str, Any],
) -> List[Dict[str, Any]]:
    variations = parsed.get('variations') or []
    if not isinstance(variations, list):
        variations = []
    achievements_pool = _collect_achievements(candidate_snapshot)
    news_pool = job_snapshot.get('company_name') and [] or []  # retained for potential future constraints

    normalized: List[Dict[str, Any]] = []
    for idx, raw in enumerate(variations):
        label = raw.get('label') or f'Variation {idx + 1}'
        var_id = slugify(label) or f'variation-{idx + 1}'
        tone = raw.get('tone') or 'balanced'

        opening = (raw.get('opening_paragraph') or '').strip()
        bodies = [
            (p or '').strip() for p in (raw.get('body_paragraphs') or []) if (p or '').strip()
        ][:MAX_PARAGRAPHS]
        closing = (raw.get('closing_paragraph') or '').strip()

        # Ensure we don't leak fabricated company or role names by lightly enforcing mentions
        # If company name is present, ensure it appears in opening paragraph
        company = job_snapshot.get('company_name') or ''
        title = job_snapshot.get('title') or ''
        if company and company.lower() not in opening.lower():
            # Soft prepend a reference
            opening = f"{company} — {opening}" if opening else company

        ach_ref = _dedupe(raw.get('achievements_referenced') or achievements_pool[:5])
        kw_used = _dedupe(raw.get('keywords_used') or job_snapshot.get('derived_keywords', []))
        news_refs = raw.get('news_citations') or []
        if not isinstance(news_refs, list):
            news_refs = []

        full_text = '\n\n'.join([p for p in [opening, *bodies, closing] if p])

        normalized.append({
            'id': var_id if var_id not in {v.get('id') for v in normalized} else f'{var_id}-{idx}',
            'label': label,
            'tone': tone,
            'opening_paragraph': opening,
            'body_paragraphs': bodies,
            'closing_paragraph': closing,
            'full_text': full_text,
            'highlights': {
                'achievements': ach_ref[:8],
                'keywords_used': kw_used[:12],
                'news_citations': news_refs[:MAX_NEWS],
            },
            'generated_at': timezone.now().isoformat(),
        })

    return normalized


def run_cover_letter_generation(
    candidate_snapshot: Dict[str, Any],
    job_snapshot: Dict[str, Any],
    research_snapshot: Dict[str, Any],
    *,
    tone: str,
    variation_count: int,
    api_key: str,
    model: str | None = None,
    length: str | None = None,
    writing_style: str | None = None,
    company_culture: str | None = None,
    industry: str | None = None,
    custom_instructions: str | None = None,
) -> Dict[str, Any]:
    logger.info("Starting cover letter generation with variation_count=%s, tone=%s", variation_count, tone)
    prompt = build_generation_prompt(
        candidate_snapshot,
        job_snapshot,
        research_snapshot,
        tone,
        variation_count,
        length=length,
        writing_style=writing_style,
        company_culture=company_culture,
        industry=industry,
        custom_instructions=custom_instructions,
    )
    raw_text = resume_ai.call_gemini_api(prompt, api_key, model=model)
    parsed = parse_payload(raw_text)
    shared_analysis = parsed.get('shared_analysis') or {}
    variations = _normalize_variations(parsed, candidate_snapshot, job_snapshot)
    if not variations:
        # Minimal fallback: synthesize a tiny generic letter using candidate and job info
        name = candidate_snapshot.get('name') or 'Candidate'
        company = job_snapshot.get('company_name') or 'the company'
        title = job_snapshot.get('title') or 'the role'
        opening = f"I’m excited to apply for the {title} role at {company}."
        bodies = [
            "My background aligns closely with your needs, and I’ve delivered measurable results in similar contexts.",
        ]
        closing = "I’d welcome the chance to discuss how I can contribute. Thank you for your time."
        variations = _normalize_variations(
            {
                'variations': [{
                    'label': 'Balanced',
                    'tone': tone,
                    'opening_paragraph': opening,
                    'body_paragraphs': bodies,
                    'closing_paragraph': closing,
                }]
            },
            candidate_snapshot,
            job_snapshot,
        )
    return {
        'shared_analysis': shared_analysis,
        'variations': variations,
        'variation_count': len(variations),
        'raw_text': raw_text,
    }


def generate_cover_letter_latex(
    candidate_name: str,
    candidate_email: str,
    candidate_phone: str,
    candidate_location: str,
    company_name: str,
    job_title: str,
    opening_paragraph: str,
    body_paragraphs: List[str],
    closing_paragraph: str,
) -> str:
    """
    Generate a LaTeX document for a cover letter.
    
    Args:
        candidate_name: Full name of the candidate
        candidate_email: Email address
        candidate_phone: Phone number
        candidate_location: City, State or location
        company_name: Name of the company
        job_title: Position title
        opening_paragraph: Opening paragraph text
        body_paragraphs: List of body paragraph texts
        closing_paragraph: Closing paragraph text
    
    Returns:
        Complete LaTeX document as a string
    """
    from datetime import date
    
    # Escape special LaTeX characters
    def latex_escape(text):
        if not text:
            return ''
        text = str(text)
        # Must escape backslash first, before other characters that produce backslashes
        text = text.replace('\\', r'\textbackslash{}')
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
            '~': r'\textasciitilde{}',
            '^': r'\textasciicircum{}',
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        return text
    
    # Escape all inputs
    name = latex_escape(candidate_name)
    email = latex_escape(candidate_email)
    phone = latex_escape(candidate_phone)
    location = latex_escape(candidate_location)
    company = latex_escape(company_name)
    title = latex_escape(job_title)
    
    # Current date
    today = date.today().strftime('%B %d, %Y')
    
    # Build the LaTeX document
    latex_lines = [
        r'\documentclass[letterpaper,11pt]{article}',
        r'\usepackage[empty]{fullpage}',
        r'\usepackage[hidelinks]{hyperref}',
        r'\usepackage{geometry}',
        r'\geometry{margin=0.75in}',
        r'\raggedright',
        r'\setlength{\tabcolsep}{0in}',
        r'\setlength{\parindent}{0pt}',
        r'\setlength{\parskip}{0.5em}',
        r'',
        r'\begin{document}',
        r'',
        f'{today}',
        r'',
        f'Hiring Manager \\\\',
        f'{company} \\\\',
        f'{title}',
        r'',
        r'\vspace{1em}',
        r'',
        f'Dear Hiring Manager,',
        r'',
    ]
    
    # Add opening paragraph
    if opening_paragraph:
        latex_lines.append(latex_escape(opening_paragraph))
        latex_lines.append('')
    
    # Add body paragraphs
    for para in body_paragraphs:
        if para and para.strip():
            latex_lines.append(latex_escape(para.strip()))
            latex_lines.append('')
    
    # Add closing paragraph
    if closing_paragraph:
        latex_lines.append(latex_escape(closing_paragraph))
        latex_lines.append('')
    
    # Add signature
    latex_lines.extend([
        r'Sincerely,',
        r'',
        f'{name}',
        r'',
        r'\end{document}',
    ])
    
    return '\n'.join(latex_lines)


def generate_cover_letter_docx(
    candidate_name: str,
    candidate_email: str,
    candidate_phone: str,
    candidate_location: str,
    company_name: str,
    job_title: str,
    opening_paragraph: str,
    body_paragraphs: List[str],
    closing_paragraph: str,
    letterhead_config: Dict[str, Any] | None = None,
) -> bytes:
    """
    Generate a Word document (.docx) for a cover letter.
    
    Args:
        candidate_name: Full name of the candidate
        candidate_email: Email address
        candidate_phone: Phone number
        candidate_location: City, State or location
        company_name: Name of the company
        job_title: Position title
        opening_paragraph: Opening paragraph text
        body_paragraphs: List of body paragraph texts
        closing_paragraph: Closing paragraph text
        letterhead_config: Optional dict with formatting preferences
    
    Returns:
        Bytes of the generated Word document
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from datetime import date
    import io
    
    # Initialize letterhead config with defaults
    config = letterhead_config or {}
    header_format = config.get('header_format', 'centered')  # 'centered', 'left', 'right'
    font_name = config.get('font_name', 'Calibri')
    font_size = config.get('font_size', 11)
    header_color = config.get('header_color', None)  # Tuple (R, G, B) or None
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add candidate header
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(candidate_name)
    header_run.bold = True
    header_run.font.size = Pt(font_size + 2)
    header_run.font.name = font_name
    if header_color:
        header_run.font.color.rgb = RGBColor(*header_color)
    
    # Set header alignment based on config
    if header_format == 'centered':
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif header_format == 'right':
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add contact info
    contact_parts = []
    if candidate_email:
        contact_parts.append(candidate_email)
    if candidate_phone:
        contact_parts.append(candidate_phone)
    if candidate_location:
        contact_parts.append(candidate_location)
    
    if contact_parts:
        contact_para = doc.add_paragraph(' | '.join(contact_parts))
        contact_para.alignment = header_para.alignment
        contact_run = contact_para.runs[0]
        contact_run.font.size = Pt(font_size - 1)
        contact_run.font.name = font_name
    
    # Add spacing
    doc.add_paragraph()
    
    # Add date
    today = date.today().strftime('%B %d, %Y')
    date_para = doc.add_paragraph(today)
    date_para.runs[0].font.size = Pt(font_size)
    date_para.runs[0].font.name = font_name
    
    # Add recipient info
    doc.add_paragraph()
    recipient_para = doc.add_paragraph('Hiring Manager')
    recipient_para.runs[0].font.size = Pt(font_size)
    recipient_para.runs[0].font.name = font_name
    
    company_para = doc.add_paragraph(company_name)
    company_para.runs[0].font.size = Pt(font_size)
    company_para.runs[0].font.name = font_name
    
    title_para = doc.add_paragraph(job_title)
    title_para.runs[0].font.size = Pt(font_size)
    title_para.runs[0].font.name = font_name
    
    # Add spacing
    doc.add_paragraph()
    
    # Add salutation
    salutation_para = doc.add_paragraph('Dear Hiring Manager,')
    salutation_para.runs[0].font.size = Pt(font_size)
    salutation_para.runs[0].font.name = font_name
    
    # Add spacing
    doc.add_paragraph()
    
    # Add opening paragraph
    if opening_paragraph:
        opening_para = doc.add_paragraph(opening_paragraph)
        opening_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        opening_para.runs[0].font.size = Pt(font_size)
        opening_para.runs[0].font.name = font_name
        doc.add_paragraph()
    
    # Add body paragraphs
    for para_text in body_paragraphs:
        if para_text and para_text.strip():
            body_para = doc.add_paragraph(para_text.strip())
            body_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_para.runs[0].font.size = Pt(font_size)
            body_para.runs[0].font.name = font_name
            doc.add_paragraph()
    
    # Add closing paragraph
    if closing_paragraph:
        closing_para = doc.add_paragraph(closing_paragraph)
        closing_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        closing_para.runs[0].font.size = Pt(font_size)
        closing_para.runs[0].font.name = font_name
        doc.add_paragraph()
    
    # Add signature
    signature_para = doc.add_paragraph('Sincerely,')
    signature_para.runs[0].font.size = Pt(font_size)
    signature_para.runs[0].font.name = font_name
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    name_para = doc.add_paragraph(candidate_name)
    name_para.runs[0].font.size = Pt(font_size)
    name_para.runs[0].font.name = font_name
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
